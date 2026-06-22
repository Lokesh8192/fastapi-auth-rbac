from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "FastAPI_Auth_RBAC_Definitions_For_Python_Learning.docx"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Definitions for Python Learning")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_topic(doc, term, definition, project_use, python_point):
    doc.add_heading(term, level=2)
    table = doc.add_table(rows=3, cols=2)
    set_table_borders(table)
    labels = ["Definition", "In this project", "Python learning point"]
    values = [definition, project_use, python_point]
    for row_index, label in enumerate(labels):
        label_cell = table.rows[row_index].cells[0]
        value_cell = table.rows[row_index].cells[1]
        label_cell.text = label
        value_cell.text = values[row_index]
        set_cell_shading(label_cell, LIGHT_BLUE)
        for cell in (label_cell, value_cell):
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(4.95)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)
    configure_styles(doc)

    doc.add_paragraph("FastAPI Auth RBAC Definitions for Python Learning", style="Title")
    doc.add_paragraph(
        "Beginner-friendly definitions for the main topics used in the fastapi-auth-rbac project.",
        style="Subtitle",
    )

    doc.add_heading("How to Study This Document", level=1)
    add_bullets(
        doc,
        [
            "First read the definition in simple words.",
            "Then connect it to where the topic appears in the project.",
            "Finally learn the Python point behind it, because most backend concepts become easier after the Python basics are clear.",
        ],
    )

    groups = {
        "Python Basics Used in the Project": [
            (
                "Python",
                "Python is a programming language used to write readable backend code, scripts, tests, and automation.",
                "The full FastAPI project is written in Python files inside the app and tests folders.",
                "Learn variables, data types, functions, classes, imports, exceptions, and virtual environments first.",
            ),
            (
                "Variable",
                "A variable is a name that stores a value so the program can use it later.",
                "Examples include user, access_token, refresh_token, db, payload, and response.",
                "Practice assigning values with name = value and reading them later in functions.",
            ),
            (
                "Function",
                "A function is a reusable block of code that performs one task.",
                "Functions like create_user, login_user, get_current_user, and create_access_token organize the project logic.",
                "Learn the def keyword, parameters, return values, and how one function can call another function.",
            ),
            (
                "def Keyword",
                "def is the Python keyword used to define a function.",
                "Routes and services are written as functions, for example def login(...): and def create_user(...):.",
                "Syntax: def function_name(parameter): then indent the function body below it.",
            ),
            (
                "Class",
                "A class is a blueprint for creating structured objects with fields and behavior.",
                "SQLAlchemy models and Pydantic schemas are classes, such as User, RefreshToken, and UserCreate.",
                "Learn class ClassName: syntax, attributes, inheritance, and object creation.",
            ),
            (
                "Object",
                "An object is a real instance created from a class.",
                "A user record returned from the database is an object of the User model.",
                "Learn dot access like user.email, user.role, and user.hashed_password.",
            ),
            (
                "Module",
                "A module is a Python file that contains code you can import elsewhere.",
                "auth.py, user_service.py, database.py, and security.py are modules.",
                "Learn import statements and how Python files share functions and classes.",
            ),
            (
                "Package",
                "A package is a folder of Python modules, usually containing an __init__.py file.",
                "app, app.api, app.services, app.models, and app.schemas act as packages.",
                "Learn how folder structure affects imports like from app.models.user import User.",
            ),
            (
                "Import",
                "An import brings code from another module into the current file.",
                "Routes import schemas, services, database dependencies, and auth helpers.",
                "Learn from module import name and why clean imports make large projects easier to read.",
            ),
            (
                "Type Hint",
                "A type hint tells readers and tools what type of value a variable or function expects.",
                "The project uses hints like db: Session, token: str, and Mapped[int].",
                "Learn basic hints such as str, int, bool, list[str], and custom class names.",
            ),
            (
                "Exception",
                "An exception is an error condition that interrupts normal code flow.",
                "The project raises HTTPException when credentials are invalid or access is forbidden.",
                "Learn try, except, raise, and finally for safe error handling.",
            ),
            (
                "Decorator",
                "A decorator adds behavior to a function using @ syntax.",
                "FastAPI route decorators like @router.post('/login') turn Python functions into API endpoints.",
                "Learn that decorators wrap functions and are common in frameworks.",
            ),
        ],
        "FastAPI and API Topics": [
            (
                "API",
                "An API is a way for one program to communicate with another program.",
                "Frontend or API clients call this backend through endpoints like /auth/login.",
                "Learn how Python functions receive input, process it, and return output as JSON.",
            ),
            (
                "FastAPI",
                "FastAPI is a Python framework for building APIs quickly with validation and documentation.",
                "It powers the authentication, refresh token, logout, current user, and admin endpoints.",
                "Learn how FastAPI maps HTTP requests to Python functions.",
            ),
            (
                "Endpoint",
                "An endpoint is a specific URL where the API accepts a request.",
                "Examples are /auth/register, /auth/login, /auth/me, /auth/refresh, /auth/logout, and /admin/users.",
                "Think of each endpoint as one Python function exposed to outside users.",
            ),
            (
                "Route",
                "A route connects an HTTP method and path to a Python function.",
                "router.post('/login') connects POST /auth/login to the login function.",
                "Learn how decorators register functions as routes.",
            ),
            (
                "APIRouter",
                "APIRouter groups related routes into separate files.",
                "The auth router groups authentication routes and the admin router groups admin routes.",
                "Learn why splitting code into modules keeps a project organized.",
            ),
            (
                "Request",
                "A request is data sent from a client to the server.",
                "LoginRequest sends email and password to the login endpoint.",
                "Learn dictionaries and JSON because request bodies are converted into Python objects.",
            ),
            (
                "Response",
                "A response is data sent back from the server to the client.",
                "The project returns messages, user data, access tokens, and refresh tokens.",
                "Learn return statements and dictionaries because API responses often start as Python dicts.",
            ),
            (
                "HTTP Method",
                "An HTTP method describes the action being requested.",
                "POST is used for register, login, refresh, and logout; GET is used for /auth/me and /admin/users.",
                "Learn the difference between reading data with GET and sending data with POST.",
            ),
            (
                "Status Code",
                "A status code is a number that tells whether an HTTP request succeeded or failed.",
                "The project uses 201 for registration, 401 for invalid auth, 403 for forbidden admin access, and 409 for duplicate data.",
                "Learn common codes: 200, 201, 400, 401, 403, 404, 409, and 422.",
            ),
            (
                "JSON",
                "JSON is a text format for sending structured data between client and server.",
                "Request bodies and responses use JSON data such as email, password, tokens, and user fields.",
                "Learn Python dictionaries because JSON objects look very similar to dicts.",
            ),
        ],
        "Authentication and Security Topics": [
            (
                "Authentication",
                "Authentication means verifying who the user is.",
                "Login authenticates a user by checking email and password.",
                "Learn conditionals and function calls because authentication is a sequence of checks.",
            ),
            (
                "Authorization",
                "Authorization means checking what an authenticated user is allowed to do.",
                "Only users with role admin can access /admin/users.",
                "Learn if statements because authorization often depends on conditions like role == 'admin'.",
            ),
            (
                "RBAC",
                "RBAC means Role-Based Access Control, where permissions depend on a user's role.",
                "The User model has a role field and admin routes require role admin.",
                "Learn strings and comparisons because roles are usually stored as text values.",
            ),
            (
                "Role",
                "A role is a label that describes a user's permission level.",
                "Default users have role user; admin users can access admin-only endpoints.",
                "Learn object attributes like current_user.role.",
            ),
            (
                "Password Hashing",
                "Password hashing converts a password into a one-way stored value.",
                "The project stores hashed_password instead of the real password.",
                "Learn why sensitive values should not be printed, logged, or saved directly.",
            ),
            (
                "bcrypt",
                "bcrypt is a password hashing algorithm designed to be slow and safe for password storage.",
                "The project uses bcrypt through passlib.",
                "Learn that security libraries should be used instead of writing your own crypto code.",
            ),
            (
                "JWT",
                "JWT means JSON Web Token, a signed token that carries claims about a user.",
                "Access tokens include sub, email, role, and exp.",
                "Learn dictionaries because JWT payload data is built from Python dicts.",
            ),
            (
                "Access Token",
                "An access token is a short-lived token used to access protected routes.",
                "The client sends it in the Authorization header to call /auth/me or /admin/users.",
                "Learn string handling because tokens are long strings passed between functions.",
            ),
            (
                "Refresh Token",
                "A refresh token is a longer-lived token used to get a new access token.",
                "The project stores refresh tokens in the database and deletes them on logout.",
                "Learn database create/read/delete flow because refresh tokens are stored records.",
            ),
            (
                "Bearer Token",
                "A bearer token is a token sent in the Authorization header using the Bearer scheme.",
                "The header looks like Authorization: Bearer <access_token>.",
                "Learn string formatting because headers are built from strings.",
            ),
            (
                "Protected Route",
                "A protected route is an endpoint that requires a valid token before access.",
                "/auth/me and /admin/users are protected routes.",
                "Learn dependency functions because protected checks are reused with Depends.",
            ),
        ],
        "Validation and Data Modeling": [
            (
                "Pydantic",
                "Pydantic validates and converts input data using Python classes.",
                "UserCreate validates registration data and LoginRequest validates login data.",
                "Learn classes and type hints because Pydantic models are class-based.",
            ),
            (
                "Schema",
                "A schema describes the shape and rules of data.",
                "UserCreate, UserResponse, and RefreshTokenRequest define what fields are expected.",
                "Learn how schemas separate external API data from internal database models.",
            ),
            (
                "Validation",
                "Validation checks whether input data follows required rules.",
                "The project validates email format, username length, password strength, and confirm password match.",
                "Learn if statements, regular expressions, and raising ValueError.",
            ),
            (
                "BaseModel",
                "BaseModel is the Pydantic parent class used to create validation schemas.",
                "Schemas like UserCreate and LoginRequest inherit from BaseModel.",
                "Learn inheritance: class UserCreate(BaseModel): means UserCreate gets BaseModel behavior.",
            ),
            (
                "Field",
                "Field adds extra validation or metadata to a Pydantic model field.",
                "Username uses Field with min_length and max_length.",
                "Learn keyword arguments like Field(min_length=3, max_length=50).",
            ),
            (
                "field_validator",
                "field_validator validates one field in a Pydantic model.",
                "The password validator checks length, uppercase, lowercase, digit, and special character.",
                "Learn decorators and class methods because validators use @field_validator and @classmethod.",
            ),
            (
                "model_validator",
                "model_validator validates the full model after individual fields are processed.",
                "It checks whether password and confirm_password match.",
                "Learn self.attribute access because the validator reads self.password and self.confirm_password.",
            ),
        ],
        "Database Topics": [
            (
                "Database",
                "A database stores application data permanently.",
                "The project stores users and refresh tokens.",
                "Learn CRUD: create, read, update, and delete.",
            ),
            (
                "PostgreSQL",
                "PostgreSQL is a powerful relational database system.",
                "The project connects to PostgreSQL through SQLAlchemy.",
                "Learn tables, rows, columns, primary keys, and foreign keys.",
            ),
            (
                "SQLAlchemy",
                "SQLAlchemy is a Python library for working with databases.",
                "It is used to define models, create sessions, and query records.",
                "Learn how Python objects can represent database rows.",
            ),
            (
                "ORM",
                "ORM means Object Relational Mapper. It maps Python classes to database tables.",
                "User and RefreshToken classes represent database tables.",
                "Learn class attributes because ORM columns are defined as attributes.",
            ),
            (
                "Model",
                "A model is a Python class that represents a database table.",
                "User is a model for the users table and RefreshToken is a model for the refresh_tokens table.",
                "Learn how one class maps to one table structure.",
            ),
            (
                "Table",
                "A table stores one type of data in rows and columns.",
                "The users table stores user records and refresh_tokens stores token records.",
                "Learn to think of each row as one object.",
            ),
            (
                "Column",
                "A column is one field in a database table.",
                "User columns include id, username, email, hashed_password, role, is_active, and created_at.",
                "Learn attribute names because column values are accessed like user.email.",
            ),
            (
                "Primary Key",
                "A primary key uniquely identifies each row in a table.",
                "User.id and RefreshToken.id are primary keys.",
                "Learn why unique ids make records easy to find.",
            ),
            (
                "Foreign Key",
                "A foreign key connects one table to another table.",
                "RefreshToken.user_id connects refresh tokens to users.",
                "Learn relationships between records, such as one user having a refresh token.",
            ),
            (
                "Session",
                "A database session is a workspace for database operations.",
                "The project uses db: Session to query, add, commit, and delete records.",
                "Learn that sessions should be opened, used, and closed properly.",
            ),
            (
                "Query",
                "A query asks the database for data.",
                "The project queries users by email, id, and username.",
                "Learn filter conditions like User.email == login_data.email.",
            ),
            (
                "Commit",
                "Commit saves pending database changes permanently.",
                "After creating a user or refresh token, the project calls db.commit().",
                "Learn that changes are not fully saved until committed.",
            ),
            (
                "Rollback",
                "Rollback cancels pending database changes after an error.",
                "The project calls db.rollback() inside exception handling.",
                "Learn try/except because rollback is part of safe database error handling.",
            ),
            (
                "Alembic",
                "Alembic is a migration tool for SQLAlchemy projects.",
                "It manages database schema changes in the alembic folder.",
                "Learn migration commands after you understand models and tables.",
            ),
            (
                "Migration",
                "A migration is a versioned database structure change.",
                "Migrations create or update tables such as users and refresh_tokens.",
                "Learn that code changes and database changes must stay in sync.",
            ),
        ],
        "Project Architecture Topics": [
            (
                "Layered Architecture",
                "Layered architecture separates responsibilities into different folders and files.",
                "Routes, services, models, schemas, dependencies, and core utilities are separated.",
                "Learn how to split large programs into smaller modules.",
            ),
            (
                "Service Layer",
                "The service layer contains business logic away from route functions.",
                "login_user and create_user are service functions.",
                "Learn that clean functions should do one clear job.",
            ),
            (
                "Dependency Injection",
                "Dependency injection provides required objects automatically instead of creating them manually inside every function.",
                "FastAPI injects db, current_user, and current_admin using Depends.",
                "Learn function parameters because dependencies are passed into route functions.",
            ),
            (
                "Depends",
                "Depends is FastAPI's way to declare a dependency.",
                "The project uses Depends(get_db), Depends(get_current_user), and Depends(get_current_admin).",
                "Learn how one function can receive the result of another function.",
            ),
            (
                "HTTPException",
                "HTTPException is used to return an HTTP error response from FastAPI.",
                "Invalid credentials raise 401, non-admin users raise 403, duplicate records raise 409.",
                "Learn raise statements and status-code based error handling.",
            ),
            (
                "Settings",
                "Settings are configuration values used by the application.",
                "SECRET_KEY, DATABASE_URL, token expiry, and algorithm are configuration values.",
                "Learn why configuration should not be hardcoded everywhere.",
            ),
            (
                "Environment Variable",
                "An environment variable is a value stored outside source code.",
                "The .env file contains configuration values for local development.",
                "Learn that secrets should come from the environment in production.",
            ),
        ],
        "Testing Topics": [
            (
                "Testing",
                "Testing checks that the code behaves correctly.",
                "The project tests registration, login, fixtures, auth headers, refresh, logout, and RBAC.",
                "Learn assert statements because tests compare expected and actual results.",
            ),
            (
                "Pytest",
                "Pytest is a Python testing framework.",
                "The tests folder uses pytest functions and fixtures.",
                "Learn test function names starting with test_.",
            ),
            (
                "Fixture",
                "A fixture prepares reusable test data or setup.",
                "test_user, registered_user, logged_in_user, and auth_headers are fixtures.",
                "Learn how fixtures reduce repeated setup code.",
            ),
            (
                "TestClient",
                "TestClient lets tests call FastAPI endpoints without manually running the server.",
                "client.post and client.get are used in tests.",
                "Learn to inspect response.status_code and response.json().",
            ),
            (
                "Dependency Override",
                "A dependency override replaces a real dependency during tests.",
                "The project replaces get_db with override_get_db for test database sessions.",
                "Learn dictionaries because overrides are stored in app.dependency_overrides.",
            ),
            (
                "Test Database",
                "A test database is a separate database used only for automated tests.",
                "Tests avoid touching real development or production data.",
                "Learn why test isolation prevents accidental data loss.",
            ),
            (
                "Cleanup",
                "Cleanup removes test data after each test.",
                "The cleanup fixture truncates users and refresh_tokens after tests.",
                "Learn yield fixtures because setup runs before yield and cleanup runs after yield.",
            ),
            (
                "Coverage",
                "Coverage measures how much code is executed by tests.",
                "pytest-cov is listed in the project requirements.",
                "Learn that high coverage helps but does not guarantee perfect tests.",
            ),
        ],
        "Tools and Commands": [
            (
                "Uvicorn",
                "Uvicorn is an ASGI server used to run FastAPI apps.",
                "You can run the app with uvicorn app.main:app --reload.",
                "Learn command-line basics because backend projects are often run from the terminal.",
            ),
            (
                "Virtual Environment",
                "A virtual environment is an isolated Python environment for project dependencies.",
                "The project has a .venv folder.",
                "Learn why each project should have its own dependencies.",
            ),
            (
                "requirements.txt",
                "requirements.txt lists Python packages needed by the project.",
                "It includes FastAPI, SQLAlchemy, Alembic, pytest, python-jose, passlib, and psycopg2-binary.",
                "Learn pip install -r requirements.txt.",
            ),
            (
                "Git",
                "Git tracks code changes over time.",
                "The project is inside a Git repository.",
                "Learn git status, git add, git commit, and git log.",
            ),
            (
                "GitHub",
                "GitHub hosts Git repositories online.",
                "The README mentions GitHub as a tool used with the project.",
                "Learn pushing code and writing useful README documentation.",
            ),
        ],
    }

    for group, topics in groups.items():
        doc.add_heading(group, level=1)
        for topic in topics:
            add_topic(doc, *topic)

    doc.add_heading("Mini Python Practice From This Project", level=1)
    add_bullets(
        doc,
        [
            "Write a function that accepts email and password and returns True or False.",
            "Create a User class with username, email, role, and is_active attributes.",
            "Create a dictionary payload with sub, email, role, and exp keys.",
            "Use if statements to check whether role == 'admin'.",
            "Write a try/except block that catches an error and prints a clean message.",
            "Create a pytest test that asserts a function returns the expected value.",
        ],
    )

    doc.add_heading("Short Revision Line", level=1)
    doc.add_paragraph(
        "This project teaches Python functions, classes, imports, decorators, exceptions, validation, database models, API routes, authentication, authorization, and automated testing through a real FastAPI backend."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
