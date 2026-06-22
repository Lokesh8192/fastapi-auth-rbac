from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "FastAPI_Auth_RBAC_Project_Notes.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("FastAPI Auth RBAC Notes")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_label_table(doc, rows, widths=(1.875, 4.625), header=None):
    table = doc.add_table(rows=0, cols=len(widths))
    set_table_borders(table)
    if header:
        row = table.add_row()
        for idx, text in enumerate(header):
            cell = row.cells[idx]
            set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
    for left, right in rows:
        row = table.add_row()
        row.cells[0].text = left
        row.cells[1].text = right
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_three_col_table(doc, header, rows, widths=(1.5, 2.5, 2.5)):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            row.cells[idx].text = text
            for paragraph in row.cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        if i != len(lines) - 1:
            run.add_break()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("FastAPI Authentication and RBAC Project Notes")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Revision guide for the fastapi-auth-rbac project: authentication, authorization, database, migrations, and testing.")

    doc.add_heading("Quick Project Summary", level=1)
    add_label_table(
        doc,
        [
            ("Project type", "Backend API for user authentication and role-based authorization."),
            ("Main framework", "FastAPI with APIRouter, Depends, automatic validation, and OpenAPI docs."),
            ("Database", "PostgreSQL accessed through SQLAlchemy ORM sessions."),
            ("Security", "bcrypt password hashing, JWT access tokens, refresh token storage, and protected endpoints."),
            ("Testing", "Pytest, FastAPI TestClient, dependency overrides, reusable fixtures, and test database cleanup."),
        ],
    )

    doc.add_heading("Best Preparation Order", level=1)
    add_numbered(
        doc,
        [
            "FastAPI basics and routing",
            "Pydantic schemas and validation",
            "SQLAlchemy models and database sessions",
            "Password hashing with bcrypt",
            "JWT access token creation and verification",
            "Refresh token flow and logout",
            "Dependency injection for protected routes",
            "RBAC and admin-only access",
            "Alembic migrations",
            "Pytest API testing and fixtures",
            "Project architecture and interview explanation",
        ],
    )

    doc.add_heading("Core Topics", level=1)

    sections = [
        (
            "FastAPI",
            [
                "FastAPI is used to build the API endpoints.",
                "APIRouter separates authentication routes from admin routes.",
                "Depends injects common logic like database sessions and current user checks.",
                "Pydantic schemas validate request bodies and shape responses.",
            ],
            ["app/api/auth.py", "app/api/admin.py", "app/main.py"],
        ),
        (
            "Authentication",
            [
                "Authentication verifies the identity of the user.",
                "The user registers, then logs in with email and password.",
                "The server verifies the hashed password and returns tokens.",
                "The /auth/me route proves that the access token identifies a valid user.",
            ],
            ["/auth/register", "/auth/login", "/auth/me"],
        ),
        (
            "Password Hashing",
            [
                "Passwords must never be stored as plain text.",
                "The project uses passlib and bcrypt to hash passwords.",
                "During login, verify_password compares the entered password with the stored hash.",
            ],
            ["app/utils/security.py", "hash_password", "verify_password"],
        ),
        (
            "JWT Access Tokens",
            [
                "JWT means JSON Web Token.",
                "The access token carries user id, email, role, and expiry time.",
                "Protected routes expect Authorization: Bearer <access_token>.",
                "Invalid or expired tokens return 401 Unauthorized.",
            ],
            ["app/core/auth.py", "create_access_token", "decode_access_token"],
        ),
        (
            "Refresh Tokens",
            [
                "Refresh tokens are used to get a new access token without logging in again.",
                "The token is saved in the refresh_tokens table.",
                "On login, old refresh tokens for the user are removed.",
                "Logout deletes the refresh token so it cannot be used again.",
            ],
            ["/auth/refresh", "/auth/logout", "app/models/refresh_token.py"],
        ),
        (
            "Authorization and RBAC",
            [
                "Authorization checks what an authenticated user is allowed to do.",
                "RBAC means Role-Based Access Control.",
                "The User model has a role field with values like user or admin.",
                "get_current_admin allows access only when the role is admin.",
            ],
            ["/admin/users", "get_current_admin", "403 Forbidden"],
        ),
        (
            "Pydantic Validation",
            [
                "Pydantic models validate incoming JSON and response data.",
                "EmailStr validates email addresses.",
                "field_validator checks password strength.",
                "model_validator checks that password and confirm_password match.",
            ],
            ["UserCreate", "UserResponse", "LoginRequest", "RefreshTokenRequest"],
        ),
        (
            "SQLAlchemy ORM",
            [
                "SQLAlchemy maps Python classes to database tables.",
                "User and RefreshToken are ORM models.",
                "A Session is used to query, add, commit, refresh, and rollback records.",
                "Database logic is kept in service functions instead of routes.",
            ],
            ["app/models/user.py", "app/db/database.py", "app/services/user_service.py"],
        ),
        (
            "PostgreSQL",
            [
                "PostgreSQL stores users and refresh tokens.",
                "The database URL comes from project settings.",
                "Tables are created and changed using migrations in real projects.",
            ],
            ["users table", "refresh_tokens table", "DATABASE_URL"],
        ),
        (
            "Alembic Migrations",
            [
                "Alembic tracks database schema changes.",
                "Migrations are safer than manually editing production tables.",
                "Common commands create a revision and upgrade the database to the latest version.",
            ],
            ["alembic revision --autogenerate -m \"message\"", "alembic upgrade head"],
        ),
        (
            "Dependency Injection",
            [
                "Dependency injection lets FastAPI provide reusable objects automatically.",
                "get_db provides a database session.",
                "get_current_user validates the bearer token and fetches the user.",
                "get_current_admin builds on get_current_user and checks the role.",
            ],
            ["Depends(get_db)", "Depends(get_current_user)", "Depends(get_current_admin)"],
        ),
        (
            "Pytest and TestClient",
            [
                "Pytest runs automated tests.",
                "TestClient calls API endpoints without manually starting the server.",
                "Fixtures create reusable test users, login data, and auth headers.",
                "Dependency override switches the real DB session to a test DB session.",
            ],
            ["tests/conftest.py", "tests/test_auth.py", "app.dependency_overrides"],
        ),
    ]

    for heading, points, refs in sections:
        doc.add_heading(heading, level=2)
        add_bullets(doc, points)
        add_label_table(doc, [("Project references", ", ".join(refs))], widths=(1.5, 5.0))

    doc.add_heading("Important Flows", level=1)
    add_three_col_table(
        doc,
        ["Flow", "Steps", "Result"],
        [
            ("Registration", "Validate request, hash password, save user.", "New user is stored safely."),
            ("Login", "Find user, verify password, create access and refresh tokens.", "Client receives bearer token data."),
            ("Protected route", "Read bearer token, decode JWT, fetch user from DB.", "Valid user can access route."),
            ("Refresh", "Decode refresh token, check DB token, create new access token.", "User continues session."),
            ("Logout", "Find refresh token in DB and delete it.", "Refresh token can no longer be used."),
            ("Admin access", "Authenticate user, then verify role is admin.", "Admin gets users list; normal user gets 403."),
        ],
    )

    doc.add_heading("Status Codes to Remember", level=1)
    add_label_table(
        doc,
        [
            ("200", "Success."),
            ("201", "Created, used after successful registration."),
            ("401", "Unauthorized, used for invalid credentials or invalid token."),
            ("403", "Forbidden, used when logged-in user lacks permission."),
            ("404", "Resource not found, such as missing refresh token."),
            ("409", "Conflict, used for duplicate username or email."),
            ("422", "Validation error from Pydantic/FastAPI."),
        ],
        widths=(1.0, 5.5),
    )

    doc.add_heading("Common Interview Answers", level=1)
    qa = [
        (
            "What is the difference between authentication and authorization?",
            "Authentication verifies who the user is. Authorization checks what that user is allowed to access.",
        ),
        (
            "Why do we hash passwords?",
            "If the database is leaked, attackers should not see plain passwords. Hashing stores a one-way representation instead.",
        ),
        (
            "Why use refresh tokens?",
            "Access tokens should expire quickly for security. Refresh tokens allow users to continue a session without logging in repeatedly.",
        ),
        (
            "Why use dependency injection in FastAPI?",
            "It keeps route functions clean, makes shared logic reusable, and allows tests to override dependencies such as the database session.",
        ),
        (
            "How does RBAC work here?",
            "The user role is stored in the database and included in the JWT. Admin routes call get_current_admin, which rejects non-admin users with 403.",
        ),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        q = p.add_run(question)
        q.bold = True
        doc.add_paragraph(answer)

    doc.add_heading("Commands to Know", level=1)
    add_code_block(
        doc,
        [
            "uvicorn app.main:app --reload",
            "pytest -v",
            "pytest --cov",
            "alembic revision --autogenerate -m \"add table\"",
            "alembic upgrade head",
        ],
    )

    doc.add_heading("Security Improvements to Mention", level=1)
    add_bullets(
        doc,
        [
            "Store refresh tokens as hashes instead of plain token strings.",
            "Add refresh token rotation.",
            "Add rate limiting for login endpoints.",
            "Use HTTPS in production.",
            "Keep SECRET_KEY and DATABASE_URL in environment variables.",
            "Add logging and monitoring for failed login attempts.",
        ],
    )

    doc.add_heading("One-Minute Project Explanation", level=1)
    doc.add_paragraph(
        "This project is a FastAPI-based authentication and authorization API. It supports user registration, login, JWT access tokens, refresh tokens, logout, protected routes, and admin-only access through RBAC. It uses PostgreSQL with SQLAlchemy ORM, Alembic for migrations, Pydantic for validation, bcrypt for password hashing, and Pytest with TestClient for automated API testing."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
