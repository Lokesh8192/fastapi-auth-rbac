from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Python_Basics_Notes.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x21, 0x21, 0x21)
MUTED = RGBColor(0x66, 0x66, 0x66)
FILL_BLUE_GRAY = "E8EEF5"
FILL_LIGHT_GRAY = "F4F6F9"
BORDER = "D9E2F0"
CODE_FILL = "F6F8FA"
CODE_BORDER = "D8DEE4"


TOPICS = [
    {
        "title": "Importing Modules",
        "definition": "A module is a Python file or library that gives extra features to your program.",
        "explanation": "You used the random module so Python can generate random numbers.",
        "syntax": "import module_name",
        "example": "import random\nprint(random.randrange(1, 50))",
        "note": "random.randrange(1, 50) returns a random number from 1 to 49.",
    },
    {
        "title": "Variables",
        "definition": "A variable stores a value so you can use it later.",
        "explanation": "Variable names like name and age point to values such as text or numbers.",
        "syntax": "variable_name = value",
        "example": 'name = "lokesh"\nage = 29\nprint(name)\nprint(age)',
        "note": "Python decides the data type automatically from the value you assign.",
    },
    {
        "title": "Data Types and type()",
        "definition": "A data type tells Python what kind of value is stored.",
        "explanation": "type() is useful while learning because it shows whether a value is a string, integer, float, list, and so on.",
        "syntax": "type(value)",
        "example": 'name = "lokesh"\nage = 29\nprint(type(name))\nprint(type(age))',
        "note": "str means text, int means whole number, float means decimal number, and list means a collection.",
    },
    {
        "title": "Printing Output",
        "definition": "print() displays output on the screen.",
        "explanation": "It is commonly used to check values while practicing Python.",
        "syntax": "print(value)",
        "example": 'print("Hello")\nprint(name)\nprint(age)',
        "note": "Text should be inside quotes; variables are written without quotes.",
    },
    {
        "title": "Type Casting",
        "definition": "Type casting converts a value from one data type to another.",
        "explanation": "You used str(age) because Python cannot directly join text and a number with +.",
        "syntax": "int(value)\nfloat(value)\nstr(value)",
        "example": 'a = int(3.14)\nb = float(3)\nc = str(4)\nd = int("5")\nprint("your age is : " + str(age))',
        "note": "int(3.14) becomes 3, float(3) becomes 3.0, and str(4) becomes text.",
    },
    {
        "title": "String Length",
        "definition": "len() returns the number of characters in a string.",
        "explanation": "Spaces also count as characters.",
        "syntax": "len(string)",
        "example": 'name = "lokesh"\nprint(len(name))',
        "note": "The length of 'lokesh' is 6.",
    },
    {
        "title": "String Methods",
        "definition": "String methods are built-in actions used to work with text.",
        "explanation": "Methods are written after the string using a dot.",
        "syntax": "string.method()",
        "example": 'name = "  lokesh  "\nprint(name.upper())\nprint(name.lower())\nprint(name.strip())\nprint(name.replace("lokesh", "Harsha"))\nprint(name.split(","))',
        "note": "upper() changes text to capitals, lower() changes it to small letters, and strip() removes extra spaces.",
    },
    {
        "title": "Counting, Replacing, and Splitting Strings",
        "definition": "count(), replace(), and split() are useful string methods for searching and changing text.",
        "explanation": "count() counts matches, replace() changes text, and split() breaks a string into a list.",
        "syntax": "string.count(value)\nstring.replace(old, new)\nstring.split(separator)",
        "example": 'name = "lokesh"\ndob = "09-06-1997,05:00am"\nprint(name.count("lokesh"))\nprint(name.replace("lokesh", "Lokesh"))\nprint(dob.split(","))',
        "note": "dob.split(',') returns ['09-06-1997', '05:00am'].",
    },
    {
        "title": "String Slicing",
        "definition": "Slicing gets only a selected part of a string.",
        "explanation": "The start index is included, but the end index is not included.",
        "syntax": "string[start:end]",
        "example": 'name = "lokesh"\nprint(name[1:4])\nprint(name[:6])',
        "note": "Python index numbers start from 0.",
    },
    {
        "title": "f-Strings",
        "definition": "An f-string inserts variables directly inside text.",
        "explanation": "It is cleaner than joining many strings with +.",
        "syntax": 'f"text {variable}"',
        "example": 'name = "lokesh"\nage = 29\nprint(f"my name is : {name}\\nand my age is : {age}")',
        "note": "\\n means new line.",
    },
    {
        "title": "Multiple Assignment",
        "definition": "Multiple assignment assigns values to more than one variable in one line.",
        "explanation": "You can assign different values to different variables or the same value to many variables.",
        "syntax": "x, y = 4, 7\na = b = c = value",
        "example": 'x, y = 4, 7\na = b = c = ["Loki", "Harsha", "Paramesh"]\nprint(x + y)\nprint(a, b, c)',
        "note": "When a = b = c = list, all three names refer to the same list object.",
    },
    {
        "title": "Functions and global",
        "definition": "A function is a reusable block of code.",
        "explanation": "global lets a function change a variable that was created outside the function.",
        "syntax": "def function_name():\n    code",
        "example": 'name = "Loki"\n\ndef myfunc():\n    global name\n    name = "Harsha"\n    print("Hello, " + name)\n\nmyfunc()\nprint("Hello, " + name)',
        "note": "Use global carefully. Most programs are easier to understand when functions return values instead.",
    },
    {
        "title": "Membership Operator",
        "definition": "The in operator checks whether a value exists inside another value.",
        "explanation": "It returns True if the value is found and False if it is not found.",
        "syntax": "value in collection",
        "example": 'txt = "The best things in life are free!"\nprint("free" in txt)',
        "note": "This is useful for searching inside strings, lists, tuples, and other collections.",
    },
    {
        "title": "if else Conditions",
        "definition": "if else is used to make decisions in a program.",
        "explanation": "Python runs the if block when the condition is True, otherwise it runs the else block.",
        "syntax": "if condition:\n    code\nelse:\n    code",
        "example": 'txt = "The best things in life are free!"\n\nif "freez" in txt:\n    print("Yes")\nelse:\n    print("No")',
        "note": "Indentation is required in Python. The spaces show which code belongs to if or else.",
    },
    {
        "title": "Escape Characters",
        "definition": "Escape characters are special characters written with a backslash.",
        "explanation": "They let you write quotes, new lines, tabs, backslashes, and other special text inside strings.",
        "syntax": "\\n  new line\n\\t  tab\n\\\\  backslash\n\\\"  double quote\n\\'  single quote",
        "example": 'print("my name is \\"lokesh\\"")\nprint("path:C\\\\Users\\\\Lokesh\\\\hp")\nprint("path:\\nC\\\\Users\\\\Lokesh\\\\hp")',
        "note": "Use \\\\ when you want to display one backslash.",
    },
    {
        "title": "Arithmetic Operators",
        "definition": "Arithmetic operators perform mathematical calculations.",
        "explanation": "They work with numbers and return a result.",
        "syntax": "+   addition\n-   subtraction\n*   multiplication\n/   division\n%   remainder\n**  power",
        "example": "x, y = 4, 7\nprint(x + y)\nprint(x - y)\nprint(x * y)\nprint(x / y)\nprint(x % y)",
        "note": "x % y gives the remainder after division.",
    },
    {
        "title": "Assignment Operators",
        "definition": "Assignment operators store or update values in variables.",
        "explanation": "Short forms like += do the same work as writing x = x + value.",
        "syntax": "x += 5\nx -= 5\nx *= 5\nx /= 5\nx %= 5\nx **= 5",
        "example": "x = 4\nx = x + 5\nx += 5\nprint(x)",
        "note": "x += 5 means add 5 to the current value of x.",
    },
    {
        "title": "Comparison Operators",
        "definition": "Comparison operators compare two values.",
        "explanation": "They always return either True or False.",
        "syntax": "==  equal\n!=  not equal\n>   greater than\n<   less than\n>=  greater than or equal\n<=  less than or equal",
        "example": "x, y = 4, 7\nprint(x == y)\nprint(x != y)\nprint(x > y)\nprint(x < y)\nprint(x >= y)\nprint(x <= y)",
        "note": "Use == for comparison. Use = only for assignment.",
    },
]


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    element.get_or_add_pPr().append(shd)


def paragraph_bottom_border(paragraph, color=BORDER, size="6", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    set_font(r, size=9, color=MUTED)
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(text)
    r._r.append(fld_end)


def add_label_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    value_run = paragraph.add_run(text)
    set_font(value_run, size=10.5, color=INK)


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Code Block"]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    shade(paragraph._p, CODE_FILL)
    paragraph_bottom_border(paragraph, CODE_BORDER, size="4", space="2")
    run = paragraph.add_run(code)
    set_font(run, name="Consolas", size=9.5, color=RGBColor(0x24, 0x29, 0x2F))


def add_topic(doc, number, topic):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    heading.paragraph_format.keep_with_next = True
    heading.add_run(f"{number}. {topic['title']}")

    add_label_paragraph(doc, "Definition", topic["definition"])
    add_label_paragraph(doc, "Explanation", topic["explanation"])

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(2)
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run("Syntax")
    set_font(run, size=10.5, color=DARK_BLUE, bold=True)
    add_code_block(doc, topic["syntax"])

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run("Example")
    set_font(run, size=10.5, color=DARK_BLUE, bold=True)
    add_code_block(doc, topic["example"])

    add_label_paragraph(doc, "Remember", topic["note"])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(16)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.25
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9.5)
    code.paragraph_format.line_spacing = 1.15


def add_intro_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    left, right = table.rows[0].cells
    set_cell_shading(left, FILL_BLUE_GRAY)
    set_cell_shading(right, FILL_LIGHT_GRAY)

    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    r = left_p.add_run("How to study")
    set_font(r, size=11, color=DARK_BLUE, bold=True)

    right_p = right.paragraphs[0]
    right_p.paragraph_format.space_after = Pt(0)
    r = right_p.add_run(
        "Read the definition first, copy the syntax, then type the example yourself. "
        "Change one value in each example so you understand what Python is doing."
    )
    set_font(r, size=10.5, color=INK)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Python Basics Notes")
    set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("Study Guide")
    set_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("Python Basics Notes")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Definitions, syntax, examples, and explanations based on hello.py")

    add_intro_table(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Topics Covered")
    set_font(run, size=12, color=DARK_BLUE, bold=True)

    summary = doc.add_paragraph()
    summary.paragraph_format.line_spacing = 1.25
    summary.add_run(
        "Variables, data types, print(), strings, slicing, f-strings, functions, conditions, "
        "random numbers, escape characters, arithmetic operators, assignment operators, "
        "and comparison operators."
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    section_heading = doc.add_paragraph(style="Heading 1")
    section_heading.add_run("Python Basics Reference")

    for index, topic in enumerate(TOPICS, start=1):
        add_topic(doc, index, topic)
        if index in {5, 9, 13}:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    final_heading = doc.add_paragraph(style="Heading 1")
    final_heading.add_run("Quick Revision")

    revision = [
        ("Use print() to see output.", 'print("Hello")'),
        ("Use type() to check data type.", "print(type(age))"),
        ("Use str() when joining text and numbers.", 'print("Age: " + str(age))'),
        ("Use slicing to get part of text.", "name[1:4]"),
        ("Use if else to make decisions.", 'if "free" in txt:\n    print("Yes")'),
        ("Use == for comparison and = for assignment.", "x == y"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4320, 5040])
    headers = table.rows[0].cells
    headers[0].text = "Remember"
    headers[1].text = "Example"
    for cell in headers:
        set_cell_shading(cell, FILL_BLUE_GRAY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=10.5, color=DARK_BLUE, bold=True)

    for text, code in revision:
        cells = table.add_row().cells
        cells[0].text = text
        cells[1].text = code
        for run in cells[1].paragraphs[0].runs:
            set_font(run, name="Consolas", size=9.5, color=INK)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
